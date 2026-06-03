"""
DOCX 파서 — python-docx로 표 구조와 텍스트 모두 추출

DOCX는 XLSX와 PDF의 중간 성격:
- 표 구조가 명확함 (셀 단위)
- 셀 병합 정보 보유 (gridSpan, vMerge)
- 문서 흐름 (텍스트 paragraph → 표 → paragraph → 표 ...)

추출 신호:
- 표 위치 (문서 내 N번째 표)
- 셀 좌표 (row, col)
- 셀 병합 (가로 병합 gridSpan, 세로 병합 vMerge)
- 셀 배경색 (헤더 식별)
- 텍스트 굵기
"""
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path


def get_cell_background(cell):
    """셀 배경색 추출 (헤더 식별용)"""
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    if fill in (None, 'auto', '000000', 'FFFFFF'):
        return None
    return fill


def get_cell_merge_info(cell):
    """셀 병합 정보 (가로 병합 + 세로 병합)"""
    tc_pr = cell._tc.tcPr
    info = []
    if tc_pr is None:
        return info

    # 가로 병합 (gridSpan)
    grid_span = tc_pr.find(qn('w:gridSpan'))
    if grid_span is not None:
        span = grid_span.get(qn('w:val'))
        if span and int(span) > 1:
            info.append(f"GRIDSPAN={span}")

    # 세로 병합 (vMerge)
    v_merge = tc_pr.find(qn('w:vMerge'))
    if v_merge is not None:
        val = v_merge.get(qn('w:val'))
        if val == 'restart':
            info.append("VMERGE_START")
        else:
            info.append("VMERGE_CONT")

    return info


def is_cell_bold(cell):
    """셀 내 텍스트가 굵음인지"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                return True
    return False


def cell_text(cell):
    """셀의 전체 텍스트 (paragraph join)"""
    texts = []
    for p in cell.paragraphs:
        text = "".join(run.text for run in p.runs)
        if text.strip():
            texts.append(text.strip())
    return " ".join(texts) if texts else "(empty)"


def parse_table(table, table_idx):
    """단일 표를 텍스트로 변환"""
    lines = []
    lines.append(f"### Table {table_idx + 1}")
    lines.append(f"  Rows: {len(table.rows)}, Columns: {len(table.columns) if table.rows else 0}")

    # 표의 시각적 그리드를 셀별로 표시
    for r_idx, row in enumerate(table.rows):
        cells_info = []
        signals = []
        for c_idx, cell in enumerate(row.cells):
            text = cell_text(cell)
            cells_info.append(text)

            sig_parts = []
            bg = get_cell_background(cell)
            if bg:
                sig_parts.append(f"BG={bg}")
            if is_cell_bold(cell):
                sig_parts.append("BOLD")
            merge = get_cell_merge_info(cell)
            sig_parts.extend(merge)
            if sig_parts:
                col_letter = chr(ord('A') + c_idx)
                signals.append(f"{col_letter}{r_idx + 1}:{','.join(sig_parts)}")

        line = f"R{r_idx + 1:03d} | " + " | ".join(cells_info)
        if signals:
            line += f"   <<{'  ;  '.join(signals)}>>"
        lines.append(line)

    return "\n".join(lines)


def parse_docx(path):
    """DOCX 파일 전체를 텍스트 신호와 함께 추출"""
    doc = Document(path)

    sections = []
    sections.append(f"# File: {Path(path).name}")
    sections.append(f"# Document body: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    sections.append("")

    # 본문 텍스트 (표 외부)
    sections.append("## Body Text (paragraphs)")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        # 정렬·굵기 신호
        alignment = p.alignment
        is_bold = any(r.bold for r in p.runs if r.bold)
        sig = []
        if alignment is not None:
            sig.append(f"align={alignment}")
        if is_bold:
            sig.append("bold")
        sig_str = f"  [{','.join(sig)}]" if sig else ""
        sections.append(f"P{i:03d}{sig_str}: {text}")
    sections.append("")

    # 표
    sections.append("## Tables")
    for i, table in enumerate(doc.tables):
        sections.append("")
        sections.append(parse_table(table, i))

    return "\n".join(sections)


if __name__ == "__main__":
    path = "/home/claude/dataset2/I컨설팅_견적서.docx"
    out = parse_docx(path)
    out_path = Path("/home/claude/poc_round3") / "I컨설팅_견적서.parsed.txt"
    out_path.write_text(out, encoding="utf-8")
    print(f"  {Path(path).name} -> {out_path.name} ({len(out)} chars)")
